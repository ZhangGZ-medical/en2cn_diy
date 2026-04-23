# -*- coding: utf-8 -*-
"""
en2cn_docx.py — 英文 DOCX → 中文 DOCX 格式保留翻译脚本模板

使用方法：
  1. 修改下方 SRC_PATH / DST_PATH
  2. 运行 Phase 1（inspect 模式）查看文档结构，确认段落索引和表格位置
  3. 填写翻译内容（PARA_TRANSLATIONS / TABLE_TRANSLATIONS）
  4. 运行 Phase 3（translate 模式）生成中文文档
  5. 运行 Phase 4（verify 模式）验证输出

核心技术：以源文件为模板克隆，只替换文字内容，不触碰 XML 格式节点。
依赖：pip install python-docx
"""

import sys
import copy

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.oxml.ns import qn

# ══════════════════════════════════════════════════════════════
# ① 配置区：修改这里
# ══════════════════════════════════════════════════════════════

SRC_PATH = r"D:\path\to\source_english.docx"   # 源英文文件
DST_PATH = r"D:\path\to\output_chinese.docx"   # 输出中文文件

# 段落翻译映射：{段落索引: "中文内容"}
# 运行 inspect_document() 先查看段落列表，再填写
PARA_TRANSLATIONS = {
    # 示例：
    # 0: "表1 缺血性脑卒中4+4分期框架概述",
    # 1: "注：BBB，血脑屏障；MMP，基质金属蛋白酶；时间边界为近似值，个体差异可能存在。",
}

# 表格翻译映射：{表格索引: {(行, 列): "中文内容"}}
# 运行 inspect_document() 先查看表格结构，再填写
TABLE_TRANSLATIONS = {
    # 示例（第0张表格）：
    # 0: {
    #     (0, 0): "分期",
    #     (0, 1): "时间范围",
    #     (0, 2): "病理生理特征",
    #     (1, 0): "超急性期",
    #     (1, 1): "0–6小时",
    #     (1, 2): "能量代谢衰竭；早期血脑屏障破坏",
    # }
}

# ══════════════════════════════════════════════════════════════
# ② 辅助函数
# ══════════════════════════════════════════════════════════════

def replace_para_text(para, new_text: str):
    """
    替换段落文字，保留第一个 run 的全部格式（字体/大小/粗斜体等），
    清空多余 run。
    """
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def replace_cell_text(cell, new_text: str):
    """
    替换单元格文字，保留格式。
    处理单元格内第一个段落；若有多段落，只替换第一段。
    """
    para = cell.paragraphs[0]
    replace_para_text(para, new_text)


def copy_cell_format(src_cell, dst_cell):
    """
    将源单元格的 tcPr（背景色、边框、宽度等）完整复制到目标单元格。
    通常不需要手动调用——以源文件为模板时格式已继承。
    """
    src_tc = src_cell._tc
    dst_tc = dst_cell._tc
    src_tcPr = src_tc.find(qn("w:tcPr"))
    if src_tcPr is not None:
        dst_tcPr = dst_tc.find(qn("w:tcPr"))
        if dst_tcPr is not None:
            dst_tc.remove(dst_tcPr)
        dst_tc.insert(0, copy.deepcopy(src_tcPr))


# ══════════════════════════════════════════════════════════════
# ③ Phase 1：检视文档结构（inspect 模式）
# ══════════════════════════════════════════════════════════════

def inspect_document(path: str):
    """
    打印文档所有段落和表格内容，用于确认翻译目标的索引位置。
    在填写翻译内容前先运行此函数。
    """
    doc = Document(path)

    print("=" * 60)
    print("📄 段落列表")
    print("=" * 60)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i:3d}] {para.text[:100]}")

    print()
    print("=" * 60)
    print("📊 表格列表")
    print("=" * 60)
    for ti, table in enumerate(doc.tables):
        print(f"\n── 表格 {ti}（{len(table.rows)} 行 × {len(table.columns)} 列）──")
        for ri, row in enumerate(table.rows):
            row_texts = []
            seen = set()
            for ci, cell in enumerate(row.cells):
                key = id(cell._tc)
                if key not in seen:
                    seen.add(key)
                    row_texts.append(f"[{ri},{ci}] {cell.text[:30]}")
            print("  " + " | ".join(row_texts))


# ══════════════════════════════════════════════════════════════
# ④ Phase 3：执行翻译替换
# ══════════════════════════════════════════════════════════════

def translate_document(src: str, dst: str,
                       para_trans: dict,
                       table_trans: dict):
    """
    以 src 文件为模板，按翻译映射替换文字，保存到 dst。

    参数：
        src         : 源英文 DOCX 路径
        dst         : 输出中文 DOCX 路径
        para_trans  : {段落索引: "中文文字"}
        table_trans : {表格索引: {(行, 列): "中文文字"}}
    """
    doc = Document(src)

    # 替换段落
    for idx, cn_text in para_trans.items():
        if idx < len(doc.paragraphs):
            replace_para_text(doc.paragraphs[idx], cn_text)
        else:
            print(f"⚠️  段落索引 {idx} 超出范围（共 {len(doc.paragraphs)} 段）")

    # 替换表格
    for ti, cell_map in table_trans.items():
        if ti >= len(doc.tables):
            print(f"⚠️  表格索引 {ti} 超出范围（共 {len(doc.tables)} 张表）")
            continue
        table = doc.tables[ti]
        seen = {}  # 记录已处理的合并单元格
        for ri, row in enumerate(table.rows):
            col_offset = 0
            for ci, cell in enumerate(row.cells):
                key = id(cell._tc)
                if key in seen:
                    col_offset += 1
                    continue
                seen[key] = True
                real_ci = ci - col_offset
                cn_text = cell_map.get((ri, real_ci)) or cell_map.get((ri, ci))
                if cn_text is not None:
                    replace_cell_text(cell, cn_text)

    doc.save(dst)
    print(f"✅ 中文版已保存：{dst}")


# ══════════════════════════════════════════════════════════════
# ⑤ Phase 4：验证输出
# ══════════════════════════════════════════════════════════════

def verify_document(path: str):
    """打印输出文件内容，用于核对翻译是否正确。"""
    inspect_document(path)


# ══════════════════════════════════════════════════════════════
# ⑥ 主入口
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="英文 DOCX → 中文 DOCX 格式保留翻译工具"
    )
    parser.add_argument(
        "--mode",
        choices=["inspect", "translate", "verify"],
        default="inspect",
        help=(
            "inspect  : 查看文档结构（填写翻译前先运行）\n"
            "translate: 执行翻译并生成输出文件\n"
            "verify   : 验证输出文件内容"
        ),
    )
    parser.add_argument("--src", default=SRC_PATH, help="源英文文件路径")
    parser.add_argument("--dst", default=DST_PATH, help="输出中文文件路径")
    args = parser.parse_args()

    if args.mode == "inspect":
        print(f"检视源文件：{args.src}\n")
        inspect_document(args.src)

    elif args.mode == "translate":
        if not PARA_TRANSLATIONS and not TABLE_TRANSLATIONS:
            print("⚠️  翻译内容为空！请先填写 PARA_TRANSLATIONS 和 TABLE_TRANSLATIONS。")
        else:
            translate_document(args.src, args.dst, PARA_TRANSLATIONS, TABLE_TRANSLATIONS)

    elif args.mode == "verify":
        print(f"验证输出文件：{args.dst}\n")
        verify_document(args.dst)
